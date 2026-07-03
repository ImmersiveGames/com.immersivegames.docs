from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Exemplos" / "GDD PEGA.docx"


def paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//a:blip"))


def main() -> None:
    doc = Document(DOCX)
    styles = Counter()
    headings = []
    image_paragraphs = []
    nonempty = 0

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip().replace("\n", " ")
        style = paragraph.style.name if paragraph.style is not None else ""
        if text:
            nonempty += 1
        styles[style] += 1
        if style.lower().startswith("heading") or style in {"Title", "Subtitle"}:
            headings.append((index, style, text[:140]))
        if paragraph_has_image(paragraph):
            image_paragraphs.append((index, style, text[:120]))

    print(f"paragraphs={len(doc.paragraphs)} nonempty={nonempty} tables={len(doc.tables)}")
    print("\nTop styles:")
    for style, count in styles.most_common(30):
        print(f"{count:4} {style}")

    print("\nHeadings:")
    for index, style, text in headings[:300]:
        print(f"{index:04d} {style:<12} {text}")
    if len(headings) > 300:
        print(f"... {len(headings) - 300} more headings")

    print("\nImage paragraphs:")
    for index, style, text in image_paragraphs:
        print(f"{index:04d} {style:<12} {text}")


if __name__ == "__main__":
    main()
